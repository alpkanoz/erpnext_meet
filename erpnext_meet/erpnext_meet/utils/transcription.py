# Copyright (c) 2026, Alpkan Öztürk and contributors
# For license information, please see license.txt

"""
Jitsi Meeting Transcription Service
Optimized for 48-core Intel Xeon with Faster-Whisper
"""

import os
import json
import tempfile
from datetime import datetime
from typing import Optional, List, Dict, Any

import frappe
from frappe import _
from frappe.utils import now_datetime, get_site_path


class TranscriptionError(Exception):
    """Custom exception for transcription errors"""
    pass


class JitsiTranscriber:
    """
    Meeting transcription service using Faster-Whisper.
    Optimized for Intel Xeon 48-core CPU.
    """
    
    def __init__(self, settings: Optional[Any] = None):
        """
        Initialize transcriber with Meeting Settings.
        
        Args:
            settings: Meeting Settings doc or None to auto-fetch
        """
        self.settings = settings or frappe.get_single("Meeting Settings")
        self.model = None
        self.diarization_pipeline = None
        
        if not self.settings.enable_transcription:
            frappe.throw(_("Transcription is disabled in Meeting Settings"))
    
    def _load_whisper_model(self):
        """Lazy load Whisper model"""
        if self.model is not None:
            return
        
        try:
            from faster_whisper import WhisperModel
        except ImportError:
            raise TranscriptionError(
                "faster-whisper is not installed. "
                "Run: pip install faster-whisper"
            )
        
        model_size = self.settings.whisper_model or "large-v2"
        cpu_threads = self.settings.cpu_threads or 48
        
        frappe.log_error(
            f"Loading Whisper model: {model_size} with {cpu_threads} threads",
            "Transcription Info"
        )
        
        self.model = WhisperModel(
            model_size,
            device="cpu",
            compute_type="int8",
            cpu_threads=cpu_threads,
            num_workers=4
        )
    
    def _load_diarization(self):
        """Lazy load speaker diarization pipeline"""
        if self.diarization_pipeline is not None:
            return
        
        hf_token = self.settings.get_password("huggingface_token")
        if not hf_token:
            frappe.log_error(
                "HuggingFace token not configured. Speaker diarization disabled.",
                "Transcription Warning"
            )
            return
        
        try:
            from pyannote.audio import Pipeline
            self.diarization_pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                use_auth_token=hf_token
            )
        except ImportError:
            frappe.log_error(
                "pyannote.audio is not installed. Speaker diarization disabled.",
                "Transcription Warning"
            )
        except Exception as e:
            frappe.log_error(
                f"Failed to load diarization: {str(e)}",
                "Transcription Error"
            )
    
    def transcribe(self, audio_file: str) -> List[Dict[str, Any]]:
        """
        Transcribe an audio file.
        
        Args:
            audio_file: Path to audio file (WAV, MP3, etc.)
            
        Returns:
            List of transcript segments with timestamps and speaker info
        """
        if not os.path.exists(audio_file):
            raise TranscriptionError(f"Audio file not found: {audio_file}")
        
        self._load_whisper_model()
        
        start_time = datetime.now()
        language = self.settings.transcription_language or "tr"
        if language == "auto":
            language = None
        
        frappe.log_error(
            f"Starting transcription: {audio_file}, language: {language}",
            "Transcription Info"
        )
        
        # Transcribe
        segments, info = self.model.transcribe(
            audio_file,
            beam_size=5,
            word_timestamps=True,
            vad_filter=True,
            language=language
        )
        
        # Build transcript
        transcript = []
        for segment in segments:
            transcript.append({
                "start": segment.start,
                "end": segment.end,
                "text": segment.text.strip(),
                "speaker": "Speaker"
            })
        
        # Add speaker diarization if available
        self._load_diarization()
        if self.diarization_pipeline and transcript:
            try:
                diarization = self.diarization_pipeline(audio_file)
                for item in transcript:
                    for turn, _, speaker_label in diarization.itertracks(yield_label=True):
                        if turn.start <= item["start"] <= turn.end:
                            item["speaker"] = speaker_label
                            break
            except Exception as e:
                frappe.log_error(
                    f"Diarization failed: {str(e)}",
                    "Transcription Warning"
                )
        
        elapsed = (datetime.now() - start_time).total_seconds()
        duration = transcript[-1]["end"] if transcript else 0
        speed_factor = duration / elapsed if elapsed > 0 else 0
        
        frappe.log_error(
            f"Transcription completed. Duration: {duration:.1f}s, "
            f"Processing: {elapsed:.1f}s, Speed: {speed_factor:.1f}x realtime",
            "Transcription Info"
        )
        
        return transcript
    
    def save_meeting_notes(
        self, 
        meeting_name: str, 
        transcript: List[Dict[str, Any]],
        duration: float = 0
    ) -> str:
        """
        Save transcript to Meeting Notes doctype.
        
        Args:
            meeting_name: Name of the Meeting document
            transcript: List of transcript segments
            duration: Audio duration in seconds
            
        Returns:
            Name of created Meeting Notes document
        """
        notes = frappe.new_doc("Meeting Notes")
        notes.meeting = meeting_name
        notes.raw_transcript = json.dumps(transcript, ensure_ascii=False)
        notes.transcription_status = "Completed"
        notes.duration = duration
        notes.insert(ignore_permissions=True)
        
        return notes.name


def export_to_txt(transcript: List[Dict[str, Any]], title: str = "Meeting Transcript") -> str:
    """
    Export transcript to TXT format.
    
    Returns:
        File content as string
    """
    lines = [
        f"{title}",
        "=" * 50,
        f"Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        ""
    ]
    
    for item in transcript:
        timestamp = f"[{int(item['start']//60):02d}:{int(item['start']%60):02d}]"
        speaker = item.get("speaker", "Speaker")
        text = item.get("text", "")
        lines.append(f"{timestamp} {speaker}: {text}")
    
    return "\n".join(lines)


def export_to_docx(transcript: List[Dict[str, Any]], title: str = "Meeting Transcript") -> bytes:
    """
    Export transcript to DOCX format.
    
    Returns:
        DOCX file content as bytes
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise TranscriptionError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = 1  # Center
    
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()
    
    # Content
    for item in transcript:
        para = doc.add_paragraph()
        
        # Timestamp
        timestamp = f"[{int(item['start']//60):02d}:{int(item['start']%60):02d}] "
        run1 = para.add_run(timestamp)
        run1.font.color.rgb = RGBColor(128, 128, 128)
        run1.font.size = Pt(9)
        
        # Speaker
        speaker = item.get("speaker", "Speaker")
        run2 = para.add_run(f"{speaker}: ")
        run2.bold = True
        
        # Text
        para.add_run(item.get("text", ""))
    
    # Save to bytes
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(transcript: List[Dict[str, Any]], title: str = "Meeting Transcript") -> bytes:
    """
    Export transcript to PDF format.
    
    Returns:
        PDF file content as bytes
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise TranscriptionError(
            "reportlab is not installed. Run: pip install reportlab"
        )
    
    import io
    buffer = io.BytesIO()
    
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(1*inch, height - 1*inch, title)
    
    c.setFont("Helvetica", 10)
    c.drawString(1*inch, height - 1.3*inch, 
                f"Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    
    # Content
    y = height - 2*inch
    c.setFont("Helvetica", 9)
    
    for item in transcript:
        timestamp = f"[{int(item['start']//60):02d}:{int(item['start']%60):02d}]"
        speaker = item.get("speaker", "Speaker")
        text = item.get("text", "")
        line = f"{timestamp} {speaker}: {text}"
        
        # Word wrap
        max_chars = 90
        while len(line) > max_chars:
            c.drawString(1*inch, y, line[:max_chars])
            line = "    " + line[max_chars:]
            y -= 15
            if y < 1*inch:
                c.showPage()
                c.setFont("Helvetica", 9)
                y = height - 1*inch
        
        c.drawString(1*inch, y, line)
        y -= 15
        
        if y < 1*inch:
            c.showPage()
            c.setFont("Helvetica", 9)
            y = height - 1*inch
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


# Background job function
def process_transcription(audio_file: str, meeting_name: str, delete_after: bool = True):
    """
    Background job to process transcription.
    
    Args:
        audio_file: Path to audio file
        meeting_name: Meeting document name
        delete_after: Whether to delete audio file after processing
    """
    try:
        transcriber = JitsiTranscriber()
        transcript = transcriber.transcribe(audio_file)
        
        # Calculate duration
        duration = transcript[-1]["end"] if transcript else 0
        
        # Save notes
        notes_name = transcriber.save_meeting_notes(meeting_name, transcript, duration)
        
        frappe.log_error(
            f"Transcription saved: {notes_name} for meeting {meeting_name}",
            "Transcription Success"
        )
        
        # Delete audio file if configured
        if delete_after and os.path.exists(audio_file):
            os.remove(audio_file)
            frappe.log_error(
                f"Deleted temporary audio file: {audio_file}",
                "Transcription Cleanup"
            )
        
        frappe.db.commit()
        
    except Exception as e:
        frappe.log_error(
            f"Transcription failed for {meeting_name}: {str(e)}\n{frappe.get_traceback()}",
            "Transcription Error"
        )
        
        # Mark as failed
        try:
            notes = frappe.get_all("Meeting Notes", 
                filters={"meeting": meeting_name, "transcription_status": "Processing"},
                limit=1
            )
            if notes:
                frappe.db.set_value("Meeting Notes", notes[0].name, 
                                   "transcription_status", "Failed")
                frappe.db.commit()
        except Exception:
            pass
